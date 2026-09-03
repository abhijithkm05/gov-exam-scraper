"""Generates the comprehensive 3-part Word documentation for the Gov Exam Tracker."""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)

doc = docx.Document()

# Page Setup (Normal Margins)
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Set base style font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

# Helper functions
def add_title(text, subtitle_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x29, 0x4A)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle_text)
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x29, 0x4A)

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x52, 0x99)

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_p(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    return p

def add_bullet(bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(bold_prefix)
    r1.bold = True
    r2 = p.add_run(text)

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Style paragraph background
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="18" w:space="8" w:color="1B5299"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def add_table(headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    # Header formatting
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_background(cell, "0F294A")
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Row formatting
    for r_idx, row in enumerate(rows):
        bg_color = "F7F9FB" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# -----------------------------------------------------------------------------
# DOCUMENT HEADER & METADATA
# -----------------------------------------------------------------------------
add_title("Gov Exam Scraper & Applied Tracker", "Complete Architectural Specification & 10-Year Rebuild Manual")

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta_p.add_run("System Status: Production | Author: Abhijith K M | Target: Central & Karnataka Recruitment")
meta_run.font.size = Pt(10)
meta_run.font.italic = True
meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph().paragraph_format.space_after = Pt(12)

# -----------------------------------------------------------------------------
# PART 1: THE LAYMAN'S STORY
# -----------------------------------------------------------------------------
add_h1("Part 1: The Layman's Story (The What, Why, and How in Plain English)")

add_h2("1. The Real-World Problem We Set Out to Solve")
add_p(
    "Applying for government exams in India is chaotic, exhausting, and stressful. Official notifications, "
    "eligibility criteria, exam schedules, and hall ticket links are scattered across dozens of clunky, "
    "poorly maintained government websites—including KPSC, KEA, RRB, SSC, SBI, IBPS, and various public sector portals."
)
add_p("To stay on top of everything manually, a candidate had to:")
add_bullet("Manual Portal Checking: ", "Remember to visit 25+ different websites every single week.")
add_bullet("Complex Navigation: ", "Click through confusing nested tables, PDF circulars, and blurry notice boards.")
add_bullet("Fear of Missing Deadlines: ", "Constantly worry about missing application deadlines or unannounced exam dates.")
add_bullet("Scattered Credentials: ", "Track application IDs, roll numbers, and fee receipts across random screenshots, folders, and emails.")
add_p(
    "Doing this manually consumes mental energy, causes anxiety, and introduces a critical point of failure: human forgetfulness."
)

add_h2("2. What We Decided to Do")
add_p(
    "We decided to build an automated digital assistant that works 24/7 without human intervention. "
    "Instead of manually visiting websites, a computer program wakes up every morning, scans every official board, "
    "organizes the data cleanly into your private Notion workspace, and alerts your phone via Telegram and Discord "
    "only when urgent action is required."
)
add_p("During development, we made a crucial design choice—The Two-Database Split:")
add_bullet("Engine 1 (The Hunter - General Feed): ", "Scans 25+ portals for new vacancies, filters by qualification, tracks 150+ listings, and archives expired ones.")
add_bullet("Engine 2 (The Butler - Applied Tracker): ", "Focuses strictly on your 9 applied exams, tracking your registration numbers, and scanning notice boards for timetable circulars and admit cards.")

add_h2("3. How the System Works (Step-by-Step in Everyday Terms)")
add_bullet("Every Morning at 07:47 AM IST: ", "A free cloud computer hosted by GitHub automatically boots up.")
add_bullet("Scanning Board Notices: ", "The cloud computer downloads the latest notice pages from KPSC, KEA, RRB, Central Banking (SBI/IBPS), HAL, and state boards.")
add_bullet("Intelligent Reading: ", "Standard pages are parsed instantly with HTML extractors. Messy circulars are read by an AI model (Groq Llama-3) to extract post names, vacancies, qualifications, and deadlines.")
add_bullet("Notion Sync: ", "The bot populates new vacancies in your general feed, archives expired ones, and updates your 9 applied exams with 'Last Checked: Today'.")
add_bullet("High-Priority Alerts: ", "If keywords like 'Admit Card', 'Hall Ticket', or 'Exam Date' are detected for your applied exams, an alert is dispatched to your Telegram and Discord.")
add_bullet("Zero Cost: ", "The cloud runner shuts down immediately upon completion. Monthly operating cost: ₹0.00.")

doc.add_page_break()

# -----------------------------------------------------------------------------
# PART 2: SOFTWARE ENGINEERING & ARCHITECTURE DOCUMENTATION
# -----------------------------------------------------------------------------
add_h1("Part 2: Software Engineering & Architecture Documentation")

add_h2("1. High-Level Architecture Overview")
add_code(
"""                     ┌────────────────────────────────────────────────────────┐
                     │          GitHub Actions Scheduled Cron Runner          │
                     │          Every Day at 07:47 AM IST (02:17 UTC)         │
                     └───────────────────────────┬────────────────────────────┘
                                                 │
                        Triggers Ubuntu-Latest Container Execution
                                                 │
                         ┌───────────────────────┴───────────────────────┐
                         │                                               │
                         ▼                                               ▼
         ┌───────────────────────────────┐               ┌───────────────────────────────┐
         │     Job 1: Vacancy Scraper    │               │  Job 2: Applied Exam Tracker  │
         │     `gov-exam-scraper scrape` │               │   `track_applied.py` Engine   │
         └───────────────┬───────────────┘               └───────────────┬───────────────┘
                         │                                               │
           Fetches 25+ Portal Endpoints                    Queries Applied Notion DB
                         │                                 (ID: 535459a275...37ef)
       ┌─────────────────┴─────────────────┐                             │
       ▼                                   ▼                             │
┌──────────────┐                   ┌──────────────┐                      │
│ Static HTML  │                   │ Dynamic JS   │                      │
│ (BS4/Parser) │                   │ (Playwright) │                      │
└──────┬───────┘                   └──────┬───────┘                      │
       │                                   │                             │
       └─────────────────┬─────────────────┘                             │
                         │                                               │
                         ▼                                               ▼
              ┌─────────────────────┐                     ┌─────────────────────────────┐
              │ Groq Llama-3 API    │                     │ Regex Pattern Scanner       │
              │ Unstructured Parsing│                     │ (Timetable, Admit Card, CBT)│
              └──────────┬──────────┘                     └──────────────┬──────────────┘
                         │                                               │
                         ├───────────────────────┬───────────────────────┤
                         ▼                                               ▼
         ┌───────────────────────────────┐               ┌───────────────────────────────┐
         │   Notion API Client (REST)    │               │ Multi-Channel Dispatcher      │
         │ - General Feed (150+ Entries) │               │ - Telegram Bot API (HTTP POST)│
         │ - Applied DB Timestamp Updates│               │ - Discord Webhook (JSON)      │
         └───────────────────────────────┘               └───────────────────────────────┘"""
)

add_h2("2. Repository & File System Structure")
add_code(
"""Govt_Exam_Scrapper/
├── .github/
│   └── workflows/
│       └── daily-run.yml            # Unified CI/CD workflow running daily at 07:47 AM IST
├── src/
│   └── gov_exam_scraper/
│       ├── __init__.py              # Package initialization
│       ├── cli.py                   # CLI entry point (`gov-exam-scraper`)
│       ├── fetch.py                 # Content fetcher (Requests + Playwright fallback)
│       ├── models.py                # Pydantic schemas (JobListing, ScraperSettings)
│       ├── notify.py                # Discord & Telegram notification dispatchers
│       ├── notion_sync.py           # Notion API integration for general vacancy feed
│       ├── track_applied.py         # Specialized watcher for applied exams & admit cards
│       └── parsers/
│           ├── __init__.py
│           ├── base.py              # Abstract parser interface
│           ├── bs4_parser.py        # Fast standard HTML table extractor
│           └── llm_parser.py        # Groq Llama-3 parser for unstructured HTML notices
├── pyproject.toml                   # Project metadata, dependencies, and CLI bindings
├── README.md                        # Project documentation
└── .env                             # Local environment secrets (git-ignored)"""
)

add_h2("3. Database Schemas (Notion Integration)")

add_h3("Database A: General Vacancies Feed")
add_p("Designed to aggregate 150+ active public notifications across India and Karnataka.")
add_table(
    ["Property Name", "Type", "Description / Purpose"],
    [
        ["Title", "Title", "Official recruitment post name"],
        ["Organization", "Select", "Recruitment authority (KPSC, KEA, SSC, RRB, etc.)"],
        ["Category", "Select", "CENTRAL, KARNATAKA, BANKING, DEFENCE, RAILWAYS"],
        ["Total Vacancies", "Number", "Reported open seats"],
        ["Education Qualification", "Multi-select", "Degree, Diploma, Engineering, Any"],
        ["Application Deadline", "Date", "Closing date for applications"],
        ["Status", "Select", "Active, Expired, Archived"],
        ["Official Link", "URL", "Direct portal/apply page"],
        ["PDF Notice", "URL", "Direct official notification circular PDF"],
        ["Content Hash", "Rich Text", "MD5 hash of title + organization for deduplication"]
    ]
)

add_h3("Database B: My Applied Exams Tracker (ID: 535459a2751646f4906c7c5e03f337ef)")
add_p("Dedicated exclusively to tracking your 9 active applications and official schedules.")
add_table(
    ["Property Name", "Type", "Description / Purpose"],
    [
        ["Exam Name", "Title", "Applied exam title"],
        ["Authority", "Select", "KPSC, KEA, RRB, HAL, SBI, IBPS, NICL, KFD"],
        ["Notification No", "Rich Text", "Official advertisement reference number"],
        ["Registration No", "Rich Text", "Unique registration/application ID"],
        ["Status", "Select", "Exam Scheduled, Date TBD, Admit Card Out, Completed"],
        ["Exam Date", "Date", "Officially confirmed or tentative examination date"],
        ["Exam Timeline", "Rich Text", "Phase overview (Prelims, Mains, PET/PST)"],
        ["Notice Board URL", "URL", "Official announcement board endpoint to monitor"],
        ["Admit Card URL", "URL", "Direct hall ticket download link"],
        ["Latest News", "Rich Text", "Title of newest circular detected on portal"],
        ["Last Checked", "Date", "ISO date timestamp refreshed by daily runner"]
    ]
)

add_h2("4. Execution Pipeline & Logic Flow")
add_h3("Step 1: General Recruitment Scraper (cli.py)")
add_bullet("Configuration: ", "Loads environment variables into strongly typed Pydantic models.")
add_bullet("Fetch Strategy: ", "Fetches registered portal endpoints using HTTP headers; falls back to Playwright if dynamic JavaScript rendering is detected.")
add_bullet("Parsing & AI: ", "Parses standard HTML tables with BeautifulSoup. Unstructured notice blocks are passed to Groq Llama-3 for structured JSON extraction.")
add_bullet("Deduplication & Sync: ", "Generates an MD5 content hash per vacancy. Skips existing entries, creates new pages, and archives posts where Deadline < today.")

add_h3("Step 2: Applied Exam Watcher (track_applied.py)")
add_bullet("Dynamic Query: ", "Pulls all active rows from the My Applied Exams Tracker database.")
add_bullet("Notice Scan: ", "Fetches each board's announcement page and strips non-content tags (<script>, <style>, <nav>).")
add_bullet("Regex Keyword Matcher: ", "Scans links against regex patterns: (time\\s*table|exam\\s*date|schedule|admit\\s*card|hall\\s*ticket|cbt).")
add_bullet("Sync & Alert: ", "Updates the Notion row with the latest circular and current timestamp. Dispatches instant Telegram and Discord alerts if a new date or admit card is found.")

add_h2("5. Cloud CI/CD Specification (.github/workflows/daily-run.yml)")
add_bullet("Schedule: ", "Runs everyday at 02:17 UTC = 07:47 AM IST.")
add_bullet("Workflow Dispatch: ", "Manual execution supported directly from GitHub Actions dashboard.")
add_bullet("Virtual Environment: ", "Runs on ubuntu-latest using Python 3.11 with cached dependencies.")
add_bullet("Browser Engine: ", "Installs headless Chromium via playwright install --with-deps chromium.")
add_bullet("Permissions: ", "contents: write enabled to support automated updates.")

doc.add_page_break()

# -----------------------------------------------------------------------------
# PART 3: MASTER RESOURCE DIRECTORY & REBUILD PLAYBOOK
# -----------------------------------------------------------------------------
add_h1("Part 3: Master Resource Directory & Step-by-Step Rebuild Guide")

add_h2("1. Complete Software, Language & Library Manifest")
add_table(
    ["Tool / Library", "Version / Platform", "Role in Architecture"],
    [
        ["Python", "3.11+", "Core programming language runtime"],
        ["GitHub Actions", "Ubuntu 22.04 LTS", "Serverless cloud runner for automated scheduling"],
        ["requests", ">= 2.31.0", "HTTP client for page downloads and REST API requests"],
        ["beautifulsoup4", ">= 4.12.0", "HTML parser for structured table extraction"],
        ["playwright", ">= 1.40.0", "Headless browser automation for JS-heavy portals"],
        ["pydantic", ">= 2.5.0", "Data modeling, schema validation, and type safety"],
        ["pydantic-settings", ">= 2.1.0", "Environment variable loading and configuration management"],
        ["groq", ">= 0.4.0", "SDK for Groq Llama-3 low-latency AI inference"],
        ["rich", ">= 13.7.0", "Terminal tables, colored status logs, and UI feedback"],
        ["click", ">= 8.1.0", "Command-line interface builder for CLI commands"]
    ]
)

add_h2("2. External Services & APIs")
add_table(
    ["Service Name", "Model / Protocol", "Authentication", "System Purpose"],
    [
        ["Groq Cloud", "llama3-70b-8192", "API Key (Bearer Token)", "Parses unstructured notice HTML into JSON schemas"],
        ["Notion API", "REST v2022-06-28", "Integration Token (Bearer)", "Database read/write for General & Applied feeds"],
        ["Telegram Bot API", "HTTPS POST /sendMessage", "Bot Token + Chat ID", "Direct mobile push notifications for urgent updates"],
        ["Discord", "HTTPS POST Webhook", "Webhook Secret URL", "Rich embed alerts in dedicated tracking channels"]
    ]
)

add_h2("3. Applied Exam Endpoints Directory")
add_table(
    ["Exam Name", "Authority", "Target Notice Board Endpoint", "Scheduled Date"],
    [
        ["HAL Trainees 2026", "HAL", "https://hal-india.co.in/Career_Listing.aspx", "06 Sept 2026 (2:00-4:30 PM)"],
        ["SBI Junior Associates", "SBI", "https://sbi.co.in/web/careers/current-openings", "September 2026 (Tentative)"],
        ["IBPS Customer Service", "IBPS", "https://www.ibps.in/", "October 2026 (Tentative)"],
        ["NICL 500 Assistants", "NICL", "https://nationalinsurance.nic.co.in/en/recruitment", "30 Oct 2026 (Mains)"],
        ["KPSC Gazetted Probationers", "KPSC", "https://kpsc.kar.nic.in/", "15 Nov 2026 (Prelims)"],
        ["KFD Forest Watcher", "KFD", "https://aranya.gov.in/", "Date TBD (1:20 Merit List)"],
        ["KEA VAO & Land Surveyor", "KEA", "https://cetonline.karnataka.gov.in/kea/", "Date TBD"],
        ["RRB Section Controller", "RRB", "https://www.rrbbnc.gov.in/", "Date TBD (CBT Schedule)"],
        ["RRB Junior Engineers", "RRB", "https://www.rrbbnc.gov.in/", "Date TBD (CBT Schedule)"]
    ]
)

add_h2("4. Step-by-Step Instructions to Rebuild From Scratch (10-Year Playbook)")
add_p("Follow these steps to reconstruct the entire platform from zero on any clean machine:")

add_h3("Step 1: Initialize Git and Virtual Environment")
add_code(
"""mkdir Govt_Exam_Scrapper && cd Govt_Exam_Scrapper
git init
python -m venv .venv
# On Windows:
.venv\\Scripts\\Activate.ps1
# On Linux/macOS:
source .venv/bin/activate
pip install --upgrade pip setuptools wheel"""
)

add_h3("Step 2: Install Package & Headless Chromium")
add_code(
"""pip install -e .
playwright install --with-deps chromium"""
)

add_h3("Step 3: Setup External Accounts & Generate API Credentials")
add_bullet("Notion: ", "Create an integration at notion.so/my-integrations. Create the two databases with the schemas in Part 2, and share both databases with the integration.")
add_bullet("Groq: ", "Generate an API key at console.groq.com/keys.")
add_bullet("Telegram: ", "Message @BotFather to create a bot. Get your chat_id via api.telegram.org/bot<TOKEN>/getUpdates.")
add_bullet("Discord: ", "Create a Webhook in your Discord channel settings.")

add_h3("Step 4: Configure Local .env File")
add_code(
"""NOTION_API_KEY=ntn_your_secret_here
NOTION_DATABASE_ID=your_general_vacancies_db_id
NOTION_APPLIED_DB_ID=535459a2751646f4906c7c5e03f337ef
GROQ_API_KEY=gsk_your_groq_key_here
TELEGRAM_BOT_TOKEN=your_telegram_bot_token
TELEGRAM_CHAT_ID=your_telegram_chat_id
DISCORD_WEBHOOK_URL=https://discord.com/api/webhooks/..."""
)

add_h3("Step 5: Configure Cloud Autopilot (.github/workflows/daily-run.yml)")
add_p("Set up GitHub Actions to trigger every morning at 02:17 UTC (07:47 AM IST):")
add_code(
"""name: Daily Exam Tracker & Notion Sync

on:
  schedule:
    - cron: '17 2 * * *'
  workflow_dispatch:

permissions:
  contents: write

jobs:
  scrape-and-sync:
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v4
      - uses: actions/setup-python@v5
        with:
          python-version: '3.11'
          cache: 'pip'
      - run: |
          pip install --upgrade pip
          pip install -e .
          playwright install --with-deps chromium
      - name: Step 1 - Scrape Vacancies
        env:
          GROQ_API_KEY: ${{ secrets.GROQ_API_KEY }}
          NOTION_API_KEY: ${{ secrets.NOTION_API_KEY }}
          NOTION_DATABASE_ID: ${{ secrets.NOTION_DATABASE_ID }}
          DISCORD_WEBHOOK_URL: ${{ secrets.DISCORD_WEBHOOK_URL }}
          TELEGRAM_BOT_TOKEN: ${{ secrets.TELEGRAM_BOT_TOKEN }}
          TELEGRAM_CHAT_ID: ${{ secrets.TELEGRAM_CHAT_ID }}
        run: gov-exam-scraper scrape --sync-notion --notify --archive-expired
      - name: Step 2 - Track Applied Exams
        env:
          NOTION_API_KEY: ${{ secrets.NOTION_API_KEY }}
          NOTION_APPLIED_DB_ID: ${{ secrets.NOTION_APPLIED_DB_ID }}
          DISCORD_WEBHOOK_URL: ${{ secrets.DISCORD_WEBHOOK_URL }}
          TELEGRAM_BOT_TOKEN: ${{ secrets.TELEGRAM_BOT_TOKEN }}
          TELEGRAM_CHAT_ID: ${{ secrets.TELEGRAM_CHAT_ID }}
        run: python src/gov_exam_scraper/track_applied.py"""
)

add_h3("Step 6: Deploy Repository Secrets to GitHub")
add_p(
    "In your GitHub repository, open Settings -> Secrets and variables -> Actions. "
    "Add all 7 environment secrets (NOTION_API_KEY, NOTION_DATABASE_ID, NOTION_APPLIED_DB_ID, "
    "GROQ_API_KEY, TELEGRAM_BOT_TOKEN, TELEGRAM_CHAT_ID, DISCORD_WEBHOOK_URL)."
)

add_h3("Step 7: Production Verification")
add_p(
    "Trigger the workflow manually from the Actions tab on GitHub. Confirm both Step 1 and Step 2 "
    "complete with green checkmarks, and verify that the Last Checked timestamps update in Notion."
)

output_filename = "Gov_Exam_Tracker_Master_Documentation.docx"
doc.save(output_filename)
print(f"\n✅ Successfully generated: {output_filename}")
